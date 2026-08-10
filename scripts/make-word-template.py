#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把《数据调研报告》Word 模板的第一章改造为 poi-tl 标签模板(一次性工具,留档备查)。

用法(仓库根目录,需 python-docx,如 /tmp/poivenv):
    python3 scripts/make-word-template.py [模板路径]
默认原地改写 common/src/main/resources/templates/data-survey-report.docx。

改造内容(只动第一章,其余章节原样保留):
1. 1.1 总体情况段落:示例文字替换为 {{tag}} 占位句
2. 1.2 数据库实例一览表:删除示例数据行(保留表头与合计行),
   表头「序号」格首部插入独立 run 的 {{schemas}} 标签(LoopRowTableRenderPolicy 的循环锚点,
   循环模板行 = 标签所在行的下一行,即第一行示例数据行改写的 [col] 标签行),
   合计行 4 个数据格填 {{sumTables}}/{{sumColumns}}/{{sumRows}}/{{sumSize}}

注意:全文只能有这一个 {{schemas}} 标签;循环行里不要再放(模板行渲染后被删除,残留的第二个标签会渲染失败)。
"""
import sys

from docx import Document

DEFAULT_PATH = 'common/src/main/resources/templates/data-survey-report.docx'

HEADER = ['序号', '数据库实例', '实例描述', '总表', '字段数', '数据行数', '占用空间']
LOOP_TAGS = ['[index]', '[name]', '[description]', '[tableCount]', '[columnCount]', '[totalRows]', '[sizeText]']
SUM_TAGS = ['{{sumTables}}', '{{sumColumns}}', '{{sumRows}}', '{{sumSize}}']

SUMMARY_11 = ('{{dsName}}数据源共计{{schemaCount}}个数据库实例,累计数据表总数{{tableCount}}张,'
              '其中空表总数{{emptyTableCount}}张,空表总率{{emptyTableRate}}%,累计字段{{columnCount}}个,'
              '其中空字段{{emptyColumnCount}}个,字段有值总率{{fillRate}}%,累计数据总行数{{totalRows}}行,'
              '累计数据总存储量约{{totalSize}}。')


def set_run_text_keep_style(paragraph, text):
    """整段替换为单 run 文本,保留首个 run 的字符格式"""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell_text(cell, text):
    """单元格替换为单段单 run 文本,保留原格式"""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_run_text_keep_style(cell.paragraphs[0], text)


def main():
    path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_PATH
    doc = Document(path)

    # ---- 1.1 总体情况段落 ----
    hit = False
    for p in doc.paragraphs:
        if p.text.strip().startswith('XXX数据源共计'):
            set_run_text_keep_style(p, SUMMARY_11)
            hit = True
            break
    if not hit:
        raise SystemExit('未找到 1.1 总体情况段落(模板已被改造?)')

    # ---- 1.2 数据库实例一览表(文档中第一个匹配表头的表格;第三章同名表不动) ----
    target = None
    for t in doc.tables:
        if [c.text.strip() for c in t.rows[0].cells] == HEADER:
            target = t
            break
    if target is None:
        raise SystemExit('未找到 1.2 数据库实例一览表')

    rows = target.rows
    if len(rows) < 3:
        raise SystemExit('1.2 表格行数异常(应至少有表头 + 示例行 + 合计行)')

    # 循环模板行 = 第一行示例数据行:逐格写 [col] 标签,保留原行格式
    for cell, tag in zip(rows[1].cells, LOOP_TAGS):
        set_cell_text(cell, tag)

    # 表头「序号」格首部插入独立 run 的 {{schemas}} 循环锚点(只能有这一个 {{schemas}},
    # 循环行里若再放一个,模板行被删后第二个标签渲染必炸 XmlValueDisconnectedException)
    head_p = rows[0].cells[0].paragraphs[0]
    anchor = head_p.add_run('{{schemas}}')
    head_p._p.remove(anchor._element)
    # 插到段落所有 run 之前(保留 w:pPr 在首位)
    from docx.oxml.ns import qn
    ppr = head_p._p.find(qn('w:pPr'))
    insert_at = 1 if ppr is not None else 0
    head_p._p.insert(insert_at, anchor._element)

    # 删除其余示例数据行(保留表头、循环模板行、合计行)
    for r in list(rows[2:-1]):
        r._element.getparent().remove(r._element)

    # 合计行:最后 4 个数据格填汇总标签(前 3 格可能横向合并,保持原样)
    sum_cells = target.rows[-1].cells
    for cell, tag in zip(sum_cells[-4:], SUM_TAGS):
        set_cell_text(cell, tag)

    doc.save(path)
    print('模板改造完成:', path)


if __name__ == '__main__':
    main()

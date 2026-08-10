#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""《数据调研报告》模板第二~四章 + 封面 poi-tl 标签改造(一次性工具,留档备查)。

前置:第一章已由 make-word-template.py 改造。本脚本在现模板基础上原地改写:
  python3 scripts/make-word-template-ch234.py [模板路径]

改造内容:
- 封面:报告名/输出时间/数据库类型 → {{reportName}}/{{outputDate}}/{{dbType}}
- 2.1 表:LoopRow 行循环({{qualityRows}} 锚点在表头首格独立 run,[col] 循环行,{{q*}} 合计/平均行)
- 2.2:引导统计段 → {{emptyAnalysis}},删除「数据库业务分组」示例表(与 2.1 口径重复)
- 2.3-2.7:引导统计段 → 对应 {{*Analysis}};分组合并示例表删除,原位留 {{*Table}} 段落标签
  (由内核 GroupTableRenderPolicy 代码建表);2.3 质量预警框文字 → {{fillRateWarning}}
- 第三章:3.1-3.3 静态内容整段删除,章标题后留 {{tagSectionsHint}} + {{tagSections}}
  (由内核 TagSectionsRenderPolicy 按 USER 标记代码生成各节)
- 第四章:标题下首段 → {{overallAnalysis}}

注意:每个标签全文档唯一;{{qualityRows}} 等锚点必须是独立 run。
"""
import copy
import sys

from docx import Document
from docx.oxml.ns import qn

DEFAULT_PATH = 'common/src/main/resources/templates/data-survey-report.docx'

W_P = qn('w:p')
W_TBL = qn('w:tbl')


def set_para_text(paragraph, text):
    """整段替换为单 run 文本,保留首个 run 的字符格式"""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell_text(cell, text):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_para_text(cell.paragraphs[0], text)


def insert_anchor_run(cell, tag):
    """单元格首段首部插入独立 run 的循环锚点标签(保留 w:pPr 在首位)"""
    p = cell.paragraphs[0]
    anchor = p.add_run(tag)
    p._p.remove(anchor._element)
    ppr = p._p.find(qn('w:pPr'))
    p._p.insert(1 if ppr is not None else 0, anchor._element)


def new_para_after(doc, ref_element, text):
    """在指定元素后插入新段落,返回段落元素"""
    p = doc.add_paragraph(text)
    ref_element.addnext(p._element)
    return p._element


def body_blocks(doc):
    """正文顶层块(段落/表格)按文档顺序列出:(元素, 段落或None, 表格或None)"""
    body = doc.element.body
    paras = {p._p: p for p in doc.paragraphs}
    tables = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == W_P and child in paras:
            yield child, paras[child], None
        elif child.tag == W_TBL and child in tables:
            yield child, None, tables[child]


def find_para(doc, text, contains=False):
    for p in doc.paragraphs:
        t = p.text.strip()
        if (contains and t.startswith(text)) or (not contains and t == text):
            return p
    raise SystemExit('未找到段落: ' + text)


def next_table_after(doc, para):
    seen = False
    for el, p, t in body_blocks(doc):
        if p is not None and p._p is para._p:
            seen = True
        elif seen and t is not None:
            return t
    raise SystemExit('段落后未找到表格: ' + para.text[:20])


def main():
    path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_PATH
    doc = Document(path)

    # ---- 封面 ----
    set_para_text(find_para(doc, 'XXX系统数据调研报告'), '{{reportName}}')
    set_para_text(find_para(doc, '输出时间：', contains=True), '输出时间:{{outputDate}}')
    set_para_text(find_para(doc, '数据库类型：', contains=True), '数据库类型:{{dbType}}')

    # ---- 2.1 整体数据质量概况表(LoopRow,与 1.2 同模式) ----
    t21 = next_table_after(doc, find_para(doc, '2.1 整体数据质量概况'))
    loop21 = ['[name]', '[tableCount]', '[emptyCount]', '[emptyRate]',
              '[columnCount]', '[emptyColumnCount]', '[emptyColumnRate]', '[avgFillRate]']
    for cell, tag in zip(t21.rows[1].cells, loop21):
        set_cell_text(cell, tag)
    insert_anchor_run(t21.rows[0].cells[0], '{{qualityRows}}')
    for r in list(t21.rows[2:-1]):
        r._element.getparent().remove(r._element)
    sum21 = ['{{qSumTables}}', '{{qSumEmpty}}', '{{qEmptyRate}}', '{{qSumColumns}}',
             '{{qSumEmptyColumns}}', '{{qEmptyColumnRate}}', '{{qAvgFillRate}}']
    for cell, tag in zip(t21.rows[-1].cells[1:], sum21):
        set_cell_text(cell, tag)

    # ---- 2.2-2.7:引导段标签化 + 分组示例表替换为代码建表标签 ----
    # (节标题定位文本含全/半角括号差异,用前缀匹配;analysis 段落为其后第一个非空正文段)
    sections = [
        ('2.2 空表分析', '{{emptyAnalysis}}', None),                # None = 删表不留标签(2.2 不再保留表格)
        ('2.3 字段填充率分析', '{{fillRateAnalysis}}', '{{fillRateTable}}'),
        ('2.4 元数据完整性分析', '{{metadataAnalysis}}', '{{metadataTable}}'),
        ('2.5 数据类型分布', '{{typeAnalysis}}', '{{typeTable}}'),
        ('2.6 数据冗余分析', '{{redundancyAnalysis}}', '{{redundancyTable}}'),
        ('2.7大体量表分析', '{{largeTableAnalysis}}', '{{largeTable}}'),
    ]
    for title_prefix, analysis_tag, table_tag in sections:
        heading = find_para(doc, title_prefix, contains=True)
        # 引导统计段 = 标题后第一个非空、非表格的段落
        seen = False
        analysis_para = None
        for el, p, t in body_blocks(doc):
            if p is not None and p._p is heading._p:
                seen = True
            elif seen and p is not None and p.text.strip():
                analysis_para = p
                break
            elif seen and t is not None:
                break
        if analysis_para is None:
            raise SystemExit('未找到分析段: ' + title_prefix)
        set_para_text(analysis_para, analysis_tag)
        # 该节第一张示例表(2.2 的「空表的的情况如下」引导句一并清空)
        table = next_table_after(doc, heading)
        if table_tag is None:
            table._element.getparent().remove(table._element)
        else:
            new_para_after(doc, table._element, table_tag)  # 先插标签段再删表
            table._element.getparent().remove(table._element)

    # 2.2 原引导句「数据库实例业务空表的的情况如下:」清空(表格已删,引导句失去意义)
    try:
        set_para_text(find_para(doc, '数据库实例业务空表的的情况如下:'.replace(':', '：')), '')
    except SystemExit:
        set_para_text(find_para(doc, '数据库实例业务空表的的情况如下'), '')

    # ---- 2.3 质量预警框(1x1 表格,保留框体换文字) ----
    for t in doc.tables:
        if len(t.rows) == 1 and len(t.rows[0].cells) == 1 and t.rows[0].cells[0].text.strip().startswith('质量预警'):
            set_cell_text(t.rows[0].cells[0], '{{fillRateWarning}}')
            break
    else:
        raise SystemExit('未找到质量预警框')

    # ---- 第三章:删 3.1-3.3 全部内容,留标签段 ----
    ch3 = find_para(doc, '三、标签数据分析')
    ch4 = find_para(doc, '四、数据源/库整体分析说明', contains=True)
    deleting = False
    for el, p, t in list(body_blocks(doc)):
        if p is not None and p._p is ch3._p:
            deleting = True
            continue
        if p is not None and p._p is ch4._p:
            break
        if deleting:
            el.getparent().remove(el)
    # 章标题后插入两段(先 sections 后 hint,addnext 逆序)
    new_para_after(doc, ch3._p, '{{tagSections}}')
    new_para_after(doc, ch3._p, '{{tagSectionsHint}}')

    # ---- 第四章:标题后第一段 → {{overallAnalysis}} ----
    seen = False
    done = False
    for el, p, t in body_blocks(doc):
        if p is not None and p._p is ch4._p:
            seen = True
        elif seen and p is not None:
            set_para_text(p, '{{overallAnalysis}}')
            done = True
            break
    if not done:
        new_para_after(doc, ch4._p, '{{overallAnalysis}}')

    doc.save(path)
    print('模板第二~四章改造完成:', path)


if __name__ == '__main__':
    main()
